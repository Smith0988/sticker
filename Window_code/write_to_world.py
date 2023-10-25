import re
import sys
import nltk
import os
import pandas as pd
from docx import Document
import requests
from bs4 import BeautifulSoup
from google_translate import translate_with_google_translate
import win32com.client

copyright_text = "Bản quyền © 2023 Minghui.org. Mọi quyền được bảo lưu."


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def resource_path_doc(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # Kiểm tra xem chương trình có đang chạy trong môi trường phát triển hay không
        if getattr(sys, 'frozen', False):
            # Nếu đang chạy từ tệp thực thi đã xây dựng bởi PyInstaller
            base_path = sys._MEIPASS
        else:
            # Nếu đang chạy trong môi trường phát triển
            base_path = os.path.abspath(".")

        # Đường dẫn đến thư mục searching_tool trên Desktop
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        searching_tool_path = os.path.join(desktop_path, "Searching_Tool")

        # Tạo thư mục searching_tool nếu nó chưa tồn tại
        if not os.path.exists(searching_tool_path):
            os.makedirs(searching_tool_path)

        # Đường dẫn tới tài nguyên trong thư mục searching_tool
        resource_path = os.path.join(searching_tool_path, relative_path)

        return resource_path
    except Exception as e:
        print(f"Error in resource_path: {str(e)}")
        return None



csv_filename = resource_path("data\link_eng_vn_gct.csv")
csv_filename_dic = resource_path("data\dic_eng_vn_data.csv")
kv_data = resource_path("data\KV_data.csv")


def find_vietnamese_link_1(english_link):
    # Đọc dữ liệu từ file CSV
    df = pd.read_csv(csv_filename)
    match = ""
    if re.match(r'https?://', english_link):
        match = re.search(r'(en\..*?\.html)', english_link)
    if match:
        english_link_1 = match.group(1)
    else:
        english_link_1 = english_link
    # Tìm link tiếng Anh trong cột 1 sử dụng biểu thức chính quy
    pattern = re.escape(english_link_1)  # Chuyển `english_link` thành biểu thức chính quy
    matches = df[df.iloc[:, 0].str.contains(pattern, regex=True, case=False)]

    if not matches.empty:
        vietnamese_link = matches.iloc[0, 1]
        if not vietnamese_link.startswith("http://") and not vietnamese_link.startswith("https://"):
            vietnamese_link = "https:" + vietnamese_link  # Thêm schema "https:" nếu cần

        return vietnamese_link
    else:
        return "Can not find Vietnamese Link"


def get_related_link(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        article_contents_1 = soup.find('div', class_='article-body-content')
        article_contents_2 = soup.find('div', class_='articleZhengwen geo cBBlack')
        if article_contents_1:
            article_contents = article_contents_1
        else:
            article_contents = article_contents_2
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
            valid_paragraphs = []
            start_collecting = False
            related_links = []
            for paragraph in paragraphs:
                text = paragraph.get_text(strip=True)

                if "Related article" in text or "Related Article" in text or "Related Report" in text or "Related report" in text:
                    start_collecting = True

                if start_collecting:
                    if 'splitted' in paragraph.get("class", []):
                        span_sections = paragraph.find_all('span', class_='section')
                        for span in span_sections:
                            text_1 = span.get_text(strip=True)
                            valid_paragraphs.append(text_1)
                            for link in paragraph.find_all('a', href=True):
                                related_links.append(link['href'])
                    else:
                        valid_paragraphs.append(text)
                        for link in paragraph.find_all('a', href=True):
                            related_links.append(link['href'])

            if valid_paragraphs:
                article_content = "\n".join(valid_paragraphs)
                return valid_paragraphs, related_links
            else:
                return [], related_links
        else:
            return [], []
    else:
        return [], []


def find_vietnamese_link(english_link):
    first_text = "Bài liên quan:"
    all_link = []
    result_link = []
    result_link.append(first_text)
    link_fail = "Can not find Vietnamese Link"
    # Đọc dữ liệu từ file CSV
    related_content, related_link = get_related_link(english_link)
    for link in related_link:
        all_link.append(link)
    for link in all_link:
        result_link.append(find_vietnamese_link_1(link))
    return related_content, result_link, all_link


def find_vietnamese_sentence(english_sentence):
    # Đọc dữ liệu từ file CSV
    df = pd.read_csv(csv_filename_dic)

    # Tìm link tiếng Anh trong cột 1
    row = df[df.iloc[:, 0] == english_sentence]

    if not row.empty:
        vietnamese_sentence = row.iloc[0, 1]
        return vietnamese_sentence
    else:
        return []


def get_vn_article_title(url):
    try:
        # kiem tra
        if not url.startswith("http://") and not url.startswith("https://"):
            url = "https:" + url  # Thêm schema "https:" nếu cần
        # Tải nội dung của trang web
        response = requests.get(url)

        # Kiểm tra nếu yêu cầu thành công (status code 200)
        if response.status_code == 200:
            # Sử dụng BeautifulSoup để phân tích cú pháp trang web
            soup = BeautifulSoup(response.content, 'html.parser')

            # Tìm thẻ div có class là 'article-title'
            article_title_tag = soup.find('h1', class_='articleTitle cABlue')

            # Kiểm tra nếu tồn tại thẻ và lấy nội dung text của tiêu đề
            if article_title_tag:
                article_title = article_title_tag.text.strip()
                return article_title
            else:
                return "Can not find Vietnamese Link"
        else:
            return "Can not find Vietnamese Link"
    except Exception as e:
        return "Can not find Vietnamese Link"


def get_en_article_title(url):
    response = requests.get(url)
    article_title = "Can not get title from web"
    if response.status_code == 200:
        # Sử dụng BeautifulSoup để phân tích cú pháp trang web
        soup = BeautifulSoup(response.content, 'html.parser')
        # Tìm thẻ div có class
        article_title_tag = soup.find('div', class_='article-title')
        article_title_tag_1 = soup.find('h2', class_='articleTitle cABlue')
        if article_title_tag:
            article_title = article_title_tag.text.strip()
        elif article_title_tag_1:
            article_title = article_title_tag_1.text.strip()
        return article_title


def write_en_article_to_doc(url):
    # Gửi yêu cầu HTTP để lấy nội dung trang web
    response = requests.get(url)
    article_title = "Can not get artical title from web"
    if response.status_code == 200:
        # Sử dụng BeautifulSoup để phân tích cú pháp trang web
        soup = BeautifulSoup(response.content, 'html.parser')
        # Tìm thẻ div có class
        article_contents = soup.find('div', class_='article-body-content')
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
        article_title_tag = soup.find('div', class_='article-title')
        article_title_tag_1 = soup.find('h2', class_='articleTitle cABlue')
        if article_title_tag:
            article_title = article_title_tag.text.strip()
        elif article_title_tag_1:
            article_title = article_title_tag_1.text.strip()
        title_name = article_title
        article_title_tag_byline = soup.find('div', class_='article-byline')
        if not article_title_tag_byline:
            article_title_tag_byline = soup.find('div', class_='dateShare cf')
        article_title_line = article_title_tag_byline.text.strip()
        parts = article_title_line.split("|")
        if len(parts) > 1:
            article_title_line = parts[1].strip()
        else:
            article_title_line = article_title_tag_byline.text.strip()

        div_elements = soup.find_all("div", class_="translation-and-category-info")

        # Lặp qua từng thẻ div
        for div in div_elements:
            # Tìm thẻ a với lớp "cBBlue" bên trong thẻ div
            a_element = div.find("a", class_="cBBlue")

            # Kiểm tra nếu thẻ a tồn tại
            if a_element:
                link = a_element["href"]

        # Tạo tập tin Word mới
        doc = Document()
        # In đậm cho tiêu đề bài báo
        bold_run = doc.add_paragraph().add_run(article_title)
        bold_run.font.bold = True

        # In nghiêng cho dòng article_title_line
        italic_run = doc.add_paragraph().add_run(article_title_line)
        italic_run.font.italic = True

        # Thêm các đoạn văn và tiêu đề vào tập tin Word
        for paragraph in paragraphs:
            if paragraph.name == 'h3':
                run = doc.add_paragraph().add_run(paragraph.get_text())
                run.font.bold = True
            elif 'splitted' in paragraph.get("class", []):
                span_sections = paragraph.find_all('span', class_='section')
                for span in span_sections:
                    img = span.find('img')
                    if img and img.has_attr('src'):
                        image_link = img['src']
                        image_link_final = "https://en.minghui.org" + image_link
                        doc.add_paragraph(image_link_final)
                    else:
                        doc.add_paragraph(span.get_text())
            else:
                p = doc.add_paragraph()
                p.add_run(paragraph.get_text())

        italic_run = doc.add_paragraph().add_run(copyright_text)
        italic_run.font.italic = True

        doc.add_heading("", 0)
        link_en = "Bản tiếng Anh: " + url
        doc.add_paragraph(link_en)
        link_cn = "Bản tiếng Hán: " + link
        doc.add_paragraph(link_cn)

        doc.save(resource_path_doc(article_title + ".docx"))

        return title_name, link_en, link_cn

    else:
        return None


def tokenize_sentences_with_name_prefix(text):
    name_prefixes = ["Mr.", "Ms.", "No."]

    # Replace name prefixes with placeholders
    for prefix in name_prefixes:
        text = text.replace(prefix, f"{prefix}DOT")

    # Tokenize sentences
    sentences = nltk.sent_tokenize(text)

    # Replace the placeholders back to name prefixes
    for prefix in name_prefixes:
        sentences = [sentence.replace(f"{prefix}DOT", prefix) for sentence in sentences]

    return sentences


def remove_empty_paragraphs(doc):
    # Tạo danh sách mới để lưu các đoạn không trống
    non_empty_paragraphs = []

    # Duyệt qua từng đoạn trong tài liệu
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(r'^\(Minghui\.org\)', '', paragraph.text)
        # Kiểm tra xem đoạn có nội dung không
        if paragraph.text.strip():  # Kiểm tra xem đoạn có ít nhất một ký tự không phải khoảng trắng hay không
            non_empty_paragraphs.append(paragraph)

    return non_empty_paragraphs


def paragraph_execute_text(english_paragragh):
    english_text = english_paragragh.text.strip()
    english_sentence_list = tokenize_sentences_with_name_prefix(english_text)
    vietnamese_sentence_list = ""
    for english_sentence in english_sentence_list:
        vietnamese_sentence = find_vietnamese_sentence(english_sentence)
        if vietnamese_sentence:
            vietnamese_sentence_list = vietnamese_sentence_list + " " + vietnamese_sentence
        else:
            vietnamese_sentence_list = vietnamese_sentence_list + " " + translate_with_google_translate(
                english_sentence)
    vietnamese_sentence_list = vietnamese_sentence_list.strip()
    return vietnamese_sentence_list


def paragraph_text_check(english_paragragh):
    link = False
    related_text = False
    copyright_check = False
    text_en = english_paragragh.text.strip()
    if 'Related Report' in text_en or 'Related report' in text_en:
        related_text = True
    if "Related Report" in text_en or "Related report" in text_en:
        related_text = True
    if 'en.minghui.org' in text_en or "en.minghui.org" in text_en:
        link = True
    if "Bản quyền" in text_en or 'Bản quyền' in text_en:
        copyright_check = True

    return link, related_text, copyright_check


def write_related_link_to_doc(file_name, url, link_en, link_cn):
    vietnamese_text = []
    # Khởi động ứng dụng Microsoft Word
    word = win32com.client.Dispatch("Word.Application")
    # Mở tài liệu Word từ thư mục của dự án
    document_path = resource_path_doc(file_name)
    doc = word.Documents.Open(document_path)
    # Di chuyển con trỏ đến cuối tài liệu
    doc.Range().Collapse(0)  # Điểm bắt đầu
    # Thêm một dòng trống (enter) để chuyển xuống dòng mới
    doc.Range().InsertAfter('\n')
    english_text, vietnam_link, english_link = find_vietnamese_link(url)
    for link in vietnam_link:
        vietnamese_text.append(get_vn_article_title(link))
    for i in range(len(english_text)):
        if i == 0:
            # write Eng
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            doc.Range().InsertAfter('\n')
            # write vn
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnam_link[i]
            doc.Range().InsertAfter('\n')
        else:
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            # Thêm hyperlink
            hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                           Address=english_link[i - 1], SubAddress="",
                                                           TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')
            # Vietnam
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnamese_text[i]
            # Thêm hyperlink
            pattern = r"\bCan not\b"

            if not re.search(pattern, vietnamese_text[i]):
                hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                               Address=vietnam_link[i], SubAddress="",
                                                               TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = copyright_text
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()

    new_paragraph.Range.Text = link_cn
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = link_en
    doc.Range().InsertAfter('\n')

    doc.Save()
    # Đóng tài liệu
    doc.Close()
    # Đóng ứng dụng Microsoft Word
    word.Quit()


def write_related_link_to_doc_1(english_text, vietnam_link, english_link, title):
    vietnamese_text = []
    # Khởi động ứng dụng Microsoft Word
    file_name = "Related article links.docx"
    doc = Document()
    document_path = resource_path_doc(file_name)
    doc.save(document_path)
    # Khởi động ứng dụng Microsoft Word
    word = win32com.client.Dispatch("Word.Application")
    # Mở tài liệu Word từ thư mục của dự án
    doc = word.Documents.Open(document_path)
    # Di chuyển con trỏ đến cuối tài liệu
    doc.Range().Collapse(0)  # Điểm bắt đầu
    # Thêm một dòng trống (enter) để chuyển xuống dòng mới
    doc.Range().InsertAfter('\n')
    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = title
    doc.Range().InsertAfter('\n')

    # english_text, vietnam_link, english_link = find_vietnamese_link(url)
    for link in vietnam_link:
        vietnamese_text.append(get_vn_article_title(link))
    for i in range(len(english_text)):
        if i == 0:
            # write Eng
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            doc.Range().InsertAfter('\n')
            # write vn
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnam_link[i]
            doc.Range().InsertAfter('\n')
        else:
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            # Thêm hyperlink
            hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                           Address=english_link[i - 1], SubAddress="",
                                                           TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')
            # Vietnam
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnamese_text[i]
            # Thêm hyperlink
            pattern = r"\bCan not\b"

            if not re.search(pattern, vietnamese_text[i]):
                hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                               Address=vietnam_link[i], SubAddress="",
                                                               TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')
    """
    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = copyright_text
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = link_en
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = link_cn
    doc.Range().InsertAfter('\n')
    """
    doc.Save()
    # Đóng tài liệu
    doc.Close()
    # Đóng ứng dụng Microsoft Word
    word.Quit()
    return document_path


def read_paragraph_in_word(url):
    file_name, link_en, link_cn = write_en_article_to_doc(url)
    file_name_translate = file_name + "_translate"
    related_article = False

    check_done = False
    # Đọc tệp văn bản
    doc = Document(resource_path_doc(file_name + ".docx"))
    non_empty = remove_empty_paragraphs(doc)
    doc_translate = Document()
    doc_translate.add_paragraph().add_run("Created by Auto_translate tool").bold = True
    doc_translate.save(resource_path_doc(file_name_translate + ".docx"))
    # Lặp qua từng đoạn trong tệp văn bản
    for paragraph in non_empty:
        link, related_text, copyright_check = paragraph_text_check(paragraph)
        if copyright_check:
            check_done = True
            break
        if related_text:
            related_article = True
            check_done = True
            break
        doc_translate = Document(resource_path_doc(file_name_translate + ".docx"))
        if link:
            doc_translate.add_paragraph(paragraph.text.strip())
        else:
            vietnamese_paragraph = paragraph_execute_text(paragraph)
            doc_translate.add_paragraph(paragraph.text.strip())
            doc_translate.add_paragraph(vietnamese_paragraph)
        # if i == 3:
        # break
        doc_translate.save(resource_path_doc(file_name_translate + ".docx"))
    doc_translate.save(resource_path_doc(file_name_translate + ".docx"))
    if related_article:
        write_related_link_to_doc(file_name_translate + ".docx", url, link_en, link_cn)

    else:
        doc_translate = Document(resource_path_doc(file_name_translate + ".docx"))
        italic_run = doc.add_paragraph().add_run(copyright_text)
        italic_run.font.italic = True
        doc.add_heading("", 0)
        doc.add_paragraph(link_cn)
        doc.add_paragraph(link_en)
        doc_translate.save(resource_path_doc(file_name_translate + ".docx"))

    return file_name, check_done

def find_translation(english_sentence):
    english_list = []
    vietnamese_list = []
    in_text = tokenize_sentences_with_name_prefix(english_sentence)
    try:
        # Đọc dữ liệu từ file CSV vào một DataFrame
        df = pd.read_csv(kv_data)
        english_text = df.iloc[:, 0].tolist()
        # Lấy cột 1 và gán vào list2
        vietnamese_text = df.iloc[:, 1].tolist()
        # Lặp qua từng hàng trong cột 1 (tiếng Anh)
        for i in range(len(english_text)):
            if english_sentence.lower() in english_text[i].lower():
                english_list.append(english_text[i])
                vietnamese_list.append(vietnamese_text[i])
                break
        if english_list:
            return english_list, vietnamese_list, in_text
        else:
            return [], [], []

    except Exception as e:
        print(f"Đã xảy ra lỗi: {str(e)}")
        return e


def find_sentence(english_sentence):
    result_list = []  # Danh sách chứa cột 1 (tiếng Anh)
    try:
        # Đọc dữ liệu từ file CSV vào một DataFrame
        df = pd.read_csv(csv_filename_dic)
        # Lấy cột 0 và gán vào list1
        list1 = df.iloc[:, 0].tolist()

        # Lấy cột 1 và gán vào list2
        list2 = df.iloc[:, 1].tolist()
        for i in range(len(list1)):
            if english_sentence.lower() in list1[i].lower():
                result_list.append(list1[i])
                result_list.append((list2[i]))
                break
        if result_list:
            return result_list
        else:
            return []
    except Exception as e:
        print(f"Đã xảy ra lỗi: {str(e)}")
        return e


#english_sentence = "Greetings to the Dafa practitioners attending the Italian Fa conference. The Dafa disciples of the Fa rectification period all have a calling to assist Master in saving sentient beings. As such, the Dafa disciples of each region have become the hope of salvation for the people in that region. For this reason, we must cultivate ourselves well and shoulder this unprecedented,"
#text_execute(english_sentence)

# Đường dẫn đến bài báo
url = "https://en.minghui.org/html/articles/2023/7/15/210315.html"

# file_test = "my_document.docx"

# write_related_link_to_doc(file_test, url)
write_en_article_to_doc(url)
# read_paragraph_in_word(title_name ,url)
# print(get_en_article_title(url))
